from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE = Path(__file__).resolve().parent
OUT = BASE / "Rapport_semaine11_tableaux_integres.docx"


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_paragraph(doc: Document, text: str, size: int = 11, bold: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    return p


def add_figure(doc: Document, img_path: Path, caption: str):
    if not img_path.exists():
        add_paragraph(doc, f"[Image introuvable: {img_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6.4))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.italic = True


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Rapport de stage - Semaine 11\n")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"
    r2 = title.add_run("Validation CPU LUT vs FPGA IP")
    r2.font.size = Pt(13)
    r2.font.name = "Calibri"

    doc.add_paragraph()
    add_paragraph(doc, "Periode: 21 - 27 avril 2026 (a completer)")
    add_paragraph(doc, "Stagiaire: Ibrahim")
    add_paragraph(doc, "Sujet: Validation fonctionnelle et temporelle de la chaine d'acquisition GNSS")

    # Introduction
    add_heading(doc, "I. Introduction", level=1)
    add_paragraph(
        doc,
        "Cette semaine a ete consacree a la validation finale comparative entre la version "
        "CPU LUT (reference logicielle) et l'IP FPGA. Les resultats presentes ci-dessous "
        "synthesent les indicateurs les plus pertinents pour le rapport."
    )

    # Activites
    add_heading(doc, "II. Activites realisees", level=1)

    add_heading(doc, "a. Validation fonctionnelle comparative CPU vs FPGA IP", level=2)
    add_paragraph(
        doc,
        "Les deux methodes ont ete comparees sur le meme jeu de 86 signaux. "
        "Le tableau suivant resume les metriques globales."
    )
    add_figure(
        doc,
        BASE / "tableau_1_validation_globale.png",
        "Tableau 1 - Validation globale CPU LUT vs FPGA IP",
    )

    add_heading(doc, "b. Concordance des sorties entre les deux architectures", level=2)
    add_paragraph(
        doc,
        "La concordance exacte des estimations (PRN, Doppler, phase) a ete verifiee "
        "fichier par fichier. Le tableau suivant resume les accords observes."
    )
    add_figure(
        doc,
        BASE / "tableau_2_concordance.png",
        "Tableau 2 - Concordance CPU vs FPGA",
    )

    add_heading(doc, "c. Analyse temporelle et gain de performance", level=2)
    add_paragraph(
        doc,
        "L'analyse temporelle met en evidence l'acceleration apportee par l'IP FPGA "
        "par rapport a l'execution CPU LUT."
    )
    add_figure(
        doc,
        BASE / "tableau_3_temps_speedup.png",
        "Tableau 3 - Temps de calcul et facteur d'acceleration",
    )

    # Conclusion
    add_heading(doc, "III. Conclusion", level=1)
    add_paragraph(
        doc,
        "La validation confirme une coherence fonctionnelle elevee entre CPU et FPGA, "
        "avec un gain de performance temporelle majeur en faveur de l'implementation FPGA. "
        "Les tableaux presentes constituent les preuves numeriques principales a integrer "
        "dans le rapport hebdomadaire."
    )

    doc.save(str(OUT))
    print(f"DOCX genere: {OUT}")


if __name__ == "__main__":
    main()
